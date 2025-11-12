"""
Serviço de integração com Azure OpenAI.
Fornece funcionalidades de:
- Geração de embeddings para busca semântica
- Resumo de transcrições
- Chat/RAG sobre conteúdo transcrito
"""

from typing import List, Dict, Any, Optional
from openai import AzureOpenAI
from loguru import logger
from tenacity import retry, stop_after_attempt, wait_exponential
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

from app.core.config import settings


class AzureOpenAIService:
    """Cliente para Azure OpenAI Service"""

    def __init__(self):
        self.client = AzureOpenAI(
            api_key=settings.AZURE_OPENAI_KEY,
            api_version=settings.AZURE_OPENAI_API_VERSION,
            azure_endpoint=settings.AZURE_OPENAI_ENDPOINT
        )
        self.chat_deployment = settings.AZURE_OPENAI_DEPLOYMENT
        self.embedding_deployment = settings.AZURE_OPENAI_EMBEDDING_DEPLOYMENT

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def generate_embeddings(self, text: str) -> List[float]:
        """
        Gera embeddings para um texto usando text-embedding-ada-002.

        Args:
            text: Texto para gerar embeddings (max ~8000 tokens)

        Returns:
            Lista de floats representando o vetor de embedding (1536 dimensões)

        Raises:
            Exception: Se a chamada à API falhar
        """
        try:
            logger.debug(f"Gerando embeddings para texto de {len(text)} caracteres")

            response = self.client.embeddings.create(
                input=text,
                model=self.embedding_deployment
            )

            embedding = response.data[0].embedding

            logger.debug(f"Embedding gerado com {len(embedding)} dimensões")

            return embedding

        except Exception as e:
            logger.error(f"Erro ao gerar embeddings: {str(e)}")
            raise

    def generate_embeddings_batch(
        self,
        texts: List[str],
        batch_size: int = 16
    ) -> List[List[float]]:
        """
        Gera embeddings para múltiplos textos em batch.

        Args:
            texts: Lista de textos
            batch_size: Tamanho do batch (default: 16)

        Returns:
            Lista de embeddings na mesma ordem dos textos
        """
        embeddings = []

        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]

            try:
                response = self.client.embeddings.create(
                    input=batch,
                    model=self.embedding_deployment
                )

                batch_embeddings = [item.embedding for item in response.data]
                embeddings.extend(batch_embeddings)

                logger.debug(
                    f"Batch {i // batch_size + 1}: "
                    f"{len(batch_embeddings)} embeddings gerados"
                )

            except Exception as e:
                logger.error(f"Erro no batch {i // batch_size + 1}: {str(e)}")
                raise

        return embeddings

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def summarize(
        self,
        transcript: str,
        max_tokens: int = 2000,
        temperature: float = 0.3
    ) -> str:
        """
        Gera um resumo detalhado de uma transcrição.

        Args:
            transcript: Texto completo da transcrição
            max_tokens: Número máximo de tokens no resumo
            temperature: Criatividade (0.0 = determinístico, 1.0 = criativo)

        Returns:
            Resumo em linguagem natural
        """
        system_prompt = """Você é um assistente especializado em resumir transcrições de áudio.
Crie resumos DETALHADOS, informativos e bem estruturados.

Diretrizes:
- Use linguagem clara e objetiva
- Crie um resumo COMPLETO e DETALHADO, não apenas os pontos principais
- Organize em seções com títulos quando apropriado
- Inclua todos os tópicos importantes discutidos
- Mantenha a ordem cronológica quando relevante
- Use bullet points para listar itens
- Destaque decisões, ações e próximos passos
- Identifique os participantes quando mencionados
- Inclua contexto e detalhes relevantes
- Seja neutro e factual"""

        user_prompt = f"""Por favor, crie um resumo DETALHADO e COMPLETO da seguinte transcrição.
O resumo deve cobrir todos os tópicos importantes discutidos, não apenas um overview.

Transcrição:

{transcript}

Resumo Detalhado:"""

        try:
            logger.info("Gerando resumo da transcrição")

            response = self.client.chat.completions.create(
                model=self.chat_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature
            )

            summary = response.choices[0].message.content.strip()

            logger.info(f"Resumo gerado com {len(summary)} caracteres")

            return summary

        except Exception as e:
            logger.error(f"Erro ao gerar resumo: {str(e)}")
            raise

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def answer_question(
        self,
        question: str,
        context_chunks: List[str],
        chat_history: Optional[List[Dict[str, str]]] = None,
        max_tokens: int = 300,
        temperature: float = 0.7
    ) -> str:
        """
        Responde a uma pergunta baseada em chunks de contexto (RAG).

        Args:
            question: Pergunta do usuário
            context_chunks: Lista de chunks relevantes recuperados do FAISS
            chat_history: Histórico de mensagens anteriores (opcional)
            max_tokens: Número máximo de tokens na resposta
            temperature: Criatividade da resposta

        Returns:
            Resposta à pergunta baseada no contexto
        """
        # Concatenar contexto
        context = "\n\n".join([
            f"[Trecho {i+1}]\n{chunk}"
            for i, chunk in enumerate(context_chunks)
        ])

        system_prompt = """Você é um assistente que responde perguntas sobre transcrições de áudio.

Regras importantes:
- Responda APENAS com base no contexto fornecido
- Se a informação não estiver no contexto, diga "Não encontrei essa informação na transcrição"
- Seja conciso e direto
- Cite trechos relevantes quando apropriado
- Mantenha um tom profissional e amigável
- Se houver múltiplos speakers mencionados, identifique-os na resposta"""

        # Construir mensagens
        messages = [{"role": "system", "content": system_prompt}]

        # Adicionar histórico se existir
        if chat_history:
            messages.extend(chat_history)

        # Adicionar contexto e pergunta
        user_message = f"""Contexto da transcrição:
{context}

Pergunta: {question}

Resposta:"""

        messages.append({"role": "user", "content": user_message})

        try:
            logger.info(f"Respondendo pergunta: {question[:100]}...")

            response = self.client.chat.completions.create(
                model=self.chat_deployment,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature
            )

            answer = response.choices[0].message.content.strip()

            logger.info(f"Resposta gerada com {len(answer)} caracteres")

            return answer

        except Exception as e:
            logger.error(f"Erro ao responder pergunta: {str(e)}")
            raise

    def generate_title(self, transcript_preview: str, max_tokens: int = 20) -> str:
        """
        Gera um título descritivo para a transcrição.

        Args:
            transcript_preview: Primeiras linhas da transcrição
            max_tokens: Máximo de tokens no título

        Returns:
            Título sugerido
        """
        system_prompt = "Você é um assistente que cria títulos descritivos e concisos."

        user_prompt = f"""Com base no início desta transcrição, sugira um título curto e descritivo:

{transcript_preview[:500]}

Título:"""

        try:
            response = self.client.chat.completions.create(
                model=self.chat_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=max_tokens,
                temperature=0.5
            )

            title = response.choices[0].message.content.strip()
            # Remover aspas se existirem
            title = title.strip('"\'')

            return title

        except Exception as e:
            logger.error(f"Erro ao gerar título: {str(e)}")
            return "Transcrição sem título"

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    def generate_meeting_minutes(
        self,
        transcript: str,
        max_tokens: int = 3000,
        temperature: float = 0.3
    ) -> Dict[str, Any]:
        """
        Gera uma ata de reunião DETALHADA e estruturada a partir de uma transcrição.

        Args:
            transcript: Texto completo da transcrição
            max_tokens: Número máximo de tokens na ata
            temperature: Criatividade (0.0 = determinístico, 1.0 = criativo)

        Returns:
            Dict com a estrutura da ata:
            {
                "title": str,
                "summary": str,
                "topics": [{"topic": str, "discussion": str}],
                "action_items": [{"item": str, "responsible": str, "deadline": str}],
                "decisions": [str],
                "next_steps": [str]
            }
        """
        system_prompt = """Você é um assistente especializado em criar atas de reunião profissionais DETALHADAS.
Analise transcrições e gere atas estruturadas, organizadas e COMPLETAS com TODOS os detalhes relevantes.

IMPORTANTE: Sua resposta deve ser um JSON válido com a seguinte estrutura:
{
    "title": "Título descritivo da reunião",
    "summary": "Resumo executivo DETALHADO em 4-6 sentenças cobrindo os principais pontos e contexto",
    "topics": [
        {"topic": "Nome do tópico", "discussion": "Descrição DETALHADA da discussão com contexto completo, pontos levantados e conclusões"}
    ],
    "action_items": [
        {"item": "Descrição DETALHADA da ação com contexto e objetivo", "responsible": "Nome ou 'A definir'", "deadline": "Data ou 'A definir'"}
    ],
    "decisions": ["Decisão tomada com contexto e justificativa"],
    "next_steps": ["Próximo passo com detalhes de execução"]
}

Diretrizes:
- Crie atas COMPLETAS e DETALHADAS, não resumidas
- Identifique TODOS os itens de ação mencionados com descrições completas
- Extraia TODAS as decisões importantes com contexto
- Liste TODOS os próximos passos discutidos com detalhes
- Para cada tópico, inclua discussão DETALHADA com todos os pontos relevantes
- Use linguagem clara, profissional e descritiva
- Mantenha a ordem cronológica dos tópicos quando relevante
- Inclua contexto e justificativas para decisões e ações
- Identifique os participantes quando mencionados
- Se não houver informação para uma seção, use array vazio []
- Se responsável ou prazo não foram mencionados, use "A definir"
- Retorne APENAS o JSON, sem texto adicional"""

        user_prompt = f"""Por favor, gere uma ata de reunião DETALHADA e COMPLETA para a seguinte transcrição.
A ata deve incluir TODOS os tópicos discutidos, TODAS as decisões tomadas, TODOS os itens de ação e TODOS os próximos passos.

{transcript}

Ata Detalhada (JSON):"""

        try:
            logger.info("Gerando ata de reunião")

            response = self.client.chat.completions.create(
                model=self.chat_deployment,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature,
                response_format={"type": "json_object"}
            )

            content = response.choices[0].message.content.strip()

            # Parse JSON
            import json
            minutes = json.loads(content)

            logger.info(f"Ata de reunião gerada com {len(minutes.get('action_items', []))} itens de ação")

            return minutes

        except json.JSONDecodeError as e:
            logger.error(f"Erro ao parsear JSON da ata: {str(e)}")
            logger.error(f"Conteúdo recebido: {content}")
            raise ValueError("Formato inválido de ata gerada")
        except Exception as e:
            logger.error(f"Erro ao gerar ata de reunião: {str(e)}")
            raise

    def estimate_tokens(self, text: str) -> int:
        """
        Estima número de tokens em um texto.
        Aproximação: 1 token ≈ 4 caracteres em português.

        Args:
            text: Texto para estimar

        Returns:
            Número estimado de tokens
        """
        return len(text) // 4

    def generate_summary_docx(
        self,
        summary_text: str,
        filename: str,
        transcript_preview: Optional[str] = None
    ) -> BytesIO:
        """
        Gera um documento .docx formatado com o resumo da transcrição.

        Args:
            summary_text: Texto do resumo
            filename: Nome do arquivo original
            transcript_preview: Preview da transcrição (opcional)

        Returns:
            BytesIO com o documento .docx
        """
        try:
            logger.info("Gerando documento .docx formatado")

            # Criar documento
            doc = Document()

            # Configurar margens
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Título
            title = doc.add_heading('Resumo Detalhado da Transcrição', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_format = title.runs[0].font
            title_format.size = Pt(18)
            title_format.bold = True
            title_format.color.rgb = RGBColor(0, 51, 102)

            # Informações do documento
            doc.add_paragraph()
            info_table = doc.add_table(rows=2, cols=2)
            info_table.style = 'Light Grid Accent 1'

            # Arquivo
            info_table.rows[0].cells[0].text = 'Arquivo:'
            info_table.rows[0].cells[1].text = filename

            # Data
            info_table.rows[1].cells[0].text = 'Data de geração:'
            info_table.rows[1].cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M')

            # Estilizar células da tabela
            for row in info_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                row.cells[0].paragraphs[0].runs[0].font.bold = True

            # Espaço
            doc.add_paragraph()

            # Seção de Resumo
            heading = doc.add_heading('Resumo', level=1)
            heading_format = heading.runs[0].font
            heading_format.color.rgb = RGBColor(0, 51, 102)

            # Processar o texto do resumo
            # Dividir em linhas e processar formatação
            lines = summary_text.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                # Detectar se é um título (linha que termina com : ou começa com ##)
                if line.startswith('##'):
                    # Título de seção
                    title_text = line.replace('##', '').strip()
                    section_heading = doc.add_heading(title_text, level=2)
                    section_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                elif line.endswith(':') and len(line) < 80:
                    # Possível título de seção
                    p = doc.add_paragraph(line)
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                elif line.startswith('- ') or line.startswith('• '):
                    # Item de lista
                    text = line.lstrip('- •').strip()
                    p = doc.add_paragraph(text, style='List Bullet')
                    p.runs[0].font.size = Pt(11)
                elif line.startswith('*') and line.endswith('*'):
                    # Texto em itálico
                    text = line.strip('*').strip()
                    p = doc.add_paragraph(text)
                    p.runs[0].font.italic = True
                    p.runs[0].font.size = Pt(11)
                else:
                    # Parágrafo normal
                    p = doc.add_paragraph(line)
                    p.runs[0].font.size = Pt(11)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Rodapé
            doc.add_paragraph()
            doc.add_paragraph()
            footer = doc.add_paragraph('_______________________________________________')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer_text = doc.add_paragraph('Documento gerado automaticamente pelo Audia')
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text.runs[0].font.size = Pt(9)
            footer_text.runs[0].font.italic = True
            footer_text.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Salvar em BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            logger.info("Documento .docx gerado com sucesso")

            return docx_buffer

        except Exception as e:
            logger.error(f"Erro ao gerar documento .docx: {str(e)}")
            raise

    def generate_meeting_minutes_docx(
        self,
        minutes_data: Dict[str, Any],
        filename: str
    ) -> BytesIO:
        """
        Gera um documento .docx formatado com a ata de reunião.

        Args:
            minutes_data: Dicionário com os dados da ata contendo:
                - title: Título da reunião
                - summary: Resumo executivo
                - topics: Lista de tópicos [{topic, discussion}]
                - action_items: Lista de ações [{item, responsible, deadline}]
                - decisions: Lista de decisões
                - next_steps: Lista de próximos passos
            filename: Nome do arquivo original da transcrição

        Returns:
            BytesIO com o documento .docx gerado
        """
        try:
            logger.info("Gerando documento .docx da ata de reunião")

            # Criar documento
            doc = Document()

            # Configurar margens
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Título principal
            title = doc.add_heading('Ata de Reunião', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_format = title.runs[0].font
            title_format.size = Pt(20)
            title_format.bold = True
            title_format.color.rgb = RGBColor(0, 51, 102)

            # Subtítulo com título da reunião
            if minutes_data.get('title'):
                subtitle = doc.add_heading(minutes_data['title'], level=1)
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_format = subtitle.runs[0].font
                subtitle_format.size = Pt(16)
                subtitle_format.color.rgb = RGBColor(0, 102, 153)

            doc.add_paragraph()

            # Informações do documento com tabela
            info_table = doc.add_table(rows=2, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_table.rows[0].cells[0].text = 'Arquivo de origem:'
            info_table.rows[0].cells[1].text = filename
            info_table.rows[1].cells[0].text = 'Data de geração:'
            info_table.rows[1].cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M')

            # Aplicar negrito às labels
            for i in range(2):
                info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph()

            # Resumo Executivo
            if minutes_data.get('summary'):
                doc.add_heading('Resumo Executivo', level=1)
                summary_para = doc.add_paragraph(minutes_data['summary'])
                summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in summary_para.runs:
                    run.font.size = Pt(11)
                doc.add_paragraph()

            # Tópicos Discutidos
            if minutes_data.get('topics') and len(minutes_data['topics']) > 0:
                doc.add_heading('Tópicos Discutidos', level=1)

                for idx, topic in enumerate(minutes_data['topics'], 1):
                    # Título do tópico
                    topic_heading = doc.add_heading(f"{idx}. {topic.get('topic', 'Tópico sem título')}", level=2)
                    topic_heading.runs[0].font.color.rgb = RGBColor(0, 102, 153)

                    # Discussão
                    discussion_para = doc.add_paragraph(topic.get('discussion', ''))
                    discussion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in discussion_para.runs:
                        run.font.size = Pt(11)

                    doc.add_paragraph()

            # Decisões Tomadas
            if minutes_data.get('decisions') and len(minutes_data['decisions']) > 0:
                doc.add_heading('Decisões Tomadas', level=1)

                for decision in minutes_data['decisions']:
                    p = doc.add_paragraph(decision, style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        run.font.size = Pt(11)

                doc.add_paragraph()

            # Itens de Ação
            if minutes_data.get('action_items') and len(minutes_data['action_items']) > 0:
                doc.add_heading('Itens de Ação', level=1)

                # Criar tabela para itens de ação
                action_table = doc.add_table(rows=1, cols=3)
                action_table.style = 'Light Grid Accent 1'

                # Cabeçalho da tabela
                header_cells = action_table.rows[0].cells
                header_cells[0].text = 'Ação'
                header_cells[1].text = 'Responsável'
                header_cells[2].text = 'Prazo'

                # Aplicar formatação ao cabeçalho
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Adicionar itens de ação
                for action_item in minutes_data['action_items']:
                    row_cells = action_table.add_row().cells
                    row_cells[0].text = action_item.get('item', '')
                    row_cells[1].text = action_item.get('responsible', 'A definir')
                    row_cells[2].text = action_item.get('deadline', 'A definir')

                    # Formatação das células
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                doc.add_paragraph()

            # Próximos Passos
            if minutes_data.get('next_steps') and len(minutes_data['next_steps']) > 0:
                doc.add_heading('Próximos Passos', level=1)

                for step in minutes_data['next_steps']:
                    p = doc.add_paragraph(step, style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        run.font.size = Pt(11)

                doc.add_paragraph()

            # Adicionar quebra de página antes do rodapé se houver muito conteúdo
            doc.add_page_break()

            # Rodapé
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer_text = doc.add_paragraph('Ata gerada automaticamente pelo Audia')
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text.runs[0].font.size = Pt(9)
            footer_text.runs[0].font.italic = True
            footer_text.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Salvar em BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            logger.info("Documento .docx da ata gerado com sucesso")

            return docx_buffer

        except Exception as e:
            logger.error(f"Erro ao gerar documento .docx da ata: {str(e)}")
            raise


# Instância global do serviço
azure_openai_service = AzureOpenAIService()
